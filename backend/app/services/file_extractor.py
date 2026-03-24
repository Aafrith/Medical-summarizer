import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from fastapi import HTTPException, UploadFile, status

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {"pdf", "txt", "doc", "docx"}


def _extract_extension(file_name: str) -> str:
    return Path(file_name).suffix.lower().replace(".", "")


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    return content.decode("utf-8", errors="ignore")


def _extract_pdf_content(content: bytes) -> dict:
    """Extracts text, images, and tables from PDF using PyMuPDF."""
    logger.info("Starting PDF content extraction using PyMuPDF")
    doc = fitz.open(stream=content, filetype="pdf")
    
    text_content = ""
    images = []
    tables = []

    for page_index in range(len(doc)):
        page = doc[page_index]
        logger.debug(f"Processing PDF page {page_index + 1}")
        
        # Extract Text
        text_content += page.get_text() + "\n"
        
        # Extract Images
        image_list = page.get_images(full=True)
        logger.info(f"Found {len(image_list)} images on page {page_index + 1}")
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            ext = base_image["ext"]
            images.append({
                "bytes": image_bytes,
                "extension": ext,
                "name": f"page_{page_index + 1}_img_{img_index + 1}.{ext}"
            })
            logger.info(f"Extracted image {img_index + 1} from page {page_index + 1}")

        # Extract Tables
        try:
            tabs = page.find_tables()
            logger.info(f"Found {len(tabs.tables)} tables on page {page_index + 1}")
            for tab_index, tab in enumerate(tabs):
                df = tab.to_pandas()
                table_text = df.to_string(index=False)
                tables.append(table_text)
                logger.info(f"Extracted table {tab_index + 1} from page {page_index + 1}")
        except Exception as e:
            logger.warning(f"Could not extract tables from page {page_index + 1}: {e}")

    doc.close()
    logger.info(f"PDF extraction complete: {len(images)} images, {len(tables)} tables found.")
    return {
        "text": text_content.strip(),
        "images": images,
        "tables": tables
    }


def _extract_docx_text(content: bytes) -> str:
    logger.info("Extracting text from DOCX")
    doc = Document(io.BytesIO(content))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text).strip()


async def extract_content_from_upload(upload_file: UploadFile, max_size_mb: int) -> dict:
    """Main entry point for extracting content from uploaded files."""
    filename = upload_file.filename or "unknown"
    extension = _extract_extension(filename)
    logger.info(f"Processing upload: {filename} (extension: {extension})")

    if extension not in ALLOWED_EXTENSIONS:
        logger.error(f"Unsupported file format: {extension}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported format. Allowed formats: PDF, TXT, DOC, DOCX.",
        )

    content = await upload_file.read()
    file_size = len(content)
    logger.info(f"File size: {file_size} bytes")

    if file_size > max_size_mb * 1024 * 1024:
        logger.error(f"File size exceeds limit: {file_size} > {max_size_mb}MB")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File exceeds {max_size_mb} MB limit.",
        )

    result = {
        "text": "",
        "images": [],
        "tables": [],
        "extension": extension,
        "file_size": file_size
    }

    try:
        if extension == "txt":
            result["text"] = _decode_text(content)
        elif extension == "pdf":
            pdf_data = _extract_pdf_content(content)
            result.update(pdf_data)
        elif extension == "docx":
            result["text"] = _extract_docx_text(content)
        elif extension == "doc":
            result["text"] = _decode_text(content)
    except Exception as e:
        logger.exception(f"Error extracting content from {filename}: {e}")
        result["text"] = f"Error: Content could not be extracted from {filename}."

    if not result["text"].strip() and not result["images"] and not result["tables"]:
        logger.warning(f"No content extracted from {filename}")
        result["text"] = f"Document content could not be fully extracted for {filename}."

    return result

